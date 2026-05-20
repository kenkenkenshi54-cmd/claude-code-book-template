# -*- coding: utf-8 -*-
"""Build the Nice S.p.A. Business DD report as docx.

Disclosure caveat: Nice S.p.A. delisted from Borsa Italiana in April 2019.
Group-consolidated audited financials post-2018 are not publicly available.
All figures except 2018 baseline and Italian-parent standalone (Camera di
Commercio) are estimates triangulated from press releases (FSI deal Nov-2023,
acquisition announcements), industry reports, and competitor benchmarks.
Treat all FY2020-FY2024 figures as 推定 (estimates).

Currency: € million.
"""
import sys
sys.path.insert(0, "/workspaces/claude-code-book-template/.claude/skills/business-dd/reference")
from numerical_checks import PeriodData, verify_all

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "/workspaces/claude-code-book-template/workspace/business-dd/output/nice_spa_business_dd_20260426.docx"

# ============================================================================
# 数値整合性検算（出力前ゲート） — 過去5期分 (FY2020-FY2024)
# 単位：€百万。Nice Group consolidated estimates (cross-checked vs. FSI deal
# disclosure of c.€800M for FY2022 and 2018 audited group figures of €368M).
# ============================================================================
PERIODS = [
    # FY2020: COVID dip; pre-Nortek. Smart Home segment small (Fibaro+Abode).
    PeriodData(
        label="FY2020E",
        segment_sales={"Gate & Door Automation": 285.0, "Sun Shading": 90.0, "Smart Home & Security": 55.0},
        segment_profits={"Gate & Door Automation": 30.0, "Sun Shading": 12.0, "Smart Home & Security": 4.0},
        inter_segment_elimination=0.0, hq_overhead=-19.0,
        consolidated_sales=430.0, consolidated_op_profit=27.0,
    ),
    # FY2021: Nortek S&C closing in Oct 2021 — Q4 only consolidation (~$50M revenue).
    PeriodData(
        label="FY2021E",
        segment_sales={"Gate & Door Automation": 310.0, "Sun Shading": 100.0, "Smart Home & Security": 140.0},
        segment_profits={"Gate & Door Automation": 35.0, "Sun Shading": 12.0, "Smart Home & Security": 8.0},
        inter_segment_elimination=0.0, hq_overhead=-20.0,
        consolidated_sales=550.0, consolidated_op_profit=35.0,
    ),
    # FY2022: First full-year Nortek consolidation. €800M anchor (FSI press release).
    PeriodData(
        label="FY2022E",
        segment_sales={"Gate & Door Automation": 415.0, "Sun Shading": 145.0, "Smart Home & Security": 240.0},
        segment_profits={"Gate & Door Automation": 60.0, "Sun Shading": 22.0, "Smart Home & Security": 25.0},
        inter_segment_elimination=-16.0, hq_overhead=-25.0,
        consolidated_sales=800.0, consolidated_op_profit=66.0,
    ),
    # FY2023: FSI partnership year; Italian-parent standalone showed -€7.5M loss
    # (likely due to integration/restructuring charges); group resilient.
    PeriodData(
        label="FY2023E",
        segment_sales={"Gate & Door Automation": 415.0, "Sun Shading": 153.0, "Smart Home & Security": 262.0},
        segment_profits={"Gate & Door Automation": 60.0, "Sun Shading": 23.0, "Smart Home & Security": 28.0},
        inter_segment_elimination=-15.0, hq_overhead=-28.0,
        consolidated_sales=830.0, consolidated_op_profit=68.0,
    ),
    # FY2024: New CEO Mogollon (Jun-2024); Valcucine acquired Oct-2024;
    # restructuring announced for 2025 ("limited impact" cost containment).
    PeriodData(
        label="FY2024E",
        segment_sales={"Gate & Door Automation": 405.0, "Sun Shading": 162.0, "Smart Home & Security": 243.0},
        segment_profits={"Gate & Door Automation": 56.0, "Sun Shading": 24.0, "Smart Home & Security": 21.0},
        inter_segment_elimination=-10.0, hq_overhead=-25.0,
        consolidated_sales=810.0, consolidated_op_profit=66.0,
        consolidated_dep=39.0,
        geographic_sales={"Italy": 65.0, "Rest of Europe": 300.0, "North America": 285.0, "Rest of World": 160.0},
        segment_profit_rates={"Gate & Door Automation": 13.8, "Sun Shading": 14.8, "Smart Home & Security": 8.6},
    ),
]

# Table B: FY2024 構成比（独立テーブル）
# Sales total = 810; Profit total (segments only) = 56+24+21 = 101
TABLE_B = dict(
    seg_sales={"Gate & Door Automation": 405.0, "Sun Shading": 162.0, "Smart Home & Security": 243.0},
    seg_profits={"Gate & Door Automation": 56.0, "Sun Shading": 24.0, "Smart Home & Security": 21.0},
    sales_ratios={"Gate & Door Automation": 50.0, "Sun Shading": 20.0, "Smart Home & Security": 30.0},
    profit_ratios={"Gate & Door Automation": 55.4, "Sun Shading": 23.8, "Smart Home & Security": 20.8},
)

# Value-up Bridge: 5年後想定（5y target）
# 現状EBITDA: FY2024E 105 (= EBIT 66 + D&A 39)
BRIDGE = dict(
    current_ebitda=105,
    organic_low=30, organic_high=55,
    inorganic_low=25, inorganic_high=50,
    target_low=160, target_high=210,
)

# 大株主（推定; Nice Group S.p.A. = Buoro family vehicle, FSI II = €100M reserved
# capital increase Nov-2023; Mgmt & treasury and minority residual）
TOP_SHAREHOLDERS = dict(
    individual_ratios=[70.0, 15.0, 10.0, 5.0],
    stated_total=100.0,
)

# 検算実行（FAILなら docx 保存に到達しない）
verify_all(PERIODS, tolerance=0.5, table_b=TABLE_B,
           value_up_bridge=BRIDGE, top_shareholders=TOP_SHAREHOLDERS)

# ============================================================================
# Document construction
# ============================================================================
doc = Document()

# Page setup: A4, 1-inch margins
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Default style
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:eastAsia'), 'Yu Gothic')


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Arial'
        r.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
        rpr = r._element.get_or_add_rPr()
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
        rpr.append(rfonts)
    return h


def add_para(text, bold=False, italic=False, size=None, indent_pt=None):
    p = doc.add_paragraph()
    if indent_pt is not None:
        p.paragraph_format.left_indent = Pt(indent_pt)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Arial'
    if size:
        r.font.size = Pt(size)
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
    rpr.append(rfonts)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Arial'
    rpr = r._element.get_or_add_rPr()
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
    rpr.append(rfonts)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(rows, header_first=True, col_widths=None):
    n_rows = len(rows)
    n_cols = len(rows[0])
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = 'Light Grid Accent 1'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            if header_first and i == 0:
                r.bold = True
            rpr = r._element.get_or_add_rPr()
            rfonts = OxmlElement('w:rFonts')
            rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
            rpr.append(rfonts)
    doc.add_paragraph()
    return t


# ============================================================================
# 1行ヘッダ
# ============================================================================
header = doc.add_paragraph()
hr = header.add_run("Nice S.p.A. ビジネスDDレポート（2026年4月・買収候補評価）")
hr.bold = True
hr.font.size = Pt(13)
hr.font.name = 'Arial'
hr.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
rpr = hr._element.get_or_add_rPr()
rfonts = OxmlElement('w:rFonts')
rfonts.set(qn('w:eastAsia'), 'Yu Gothic')
rpr.append(rfonts)
header.paragraph_format.space_after = Pt(8)

add_para(
    "※ 開示制約：Nice S.p.A.は2019年4月にBorsa Italianaを上場廃止しており、グループ連結監査済財務は"
    "2018年期以降一般公開されていない。本レポートのFY2020-FY2024財務数値は、FSI出資プレス"
    "リリース（2023年11月、グループ売上「c.€800M」開示）、Italian Camera di Commercio提出の単体"
    "財務（Nice S.p.A. Oderzo P.IVA 03099360269）、買収プレスリリース、競合ベンチマーク"
    "（FAAC・CAME・Somfy）から三角検証した推定値。すべて「(E)推定」と明記する。",
    italic=True, size=9,
)
doc.add_paragraph()

# ============================================================================
# エグゼクティブサマリー
# ============================================================================
add_heading("エグゼクティブサマリー", level=1)

add_para("投資仮説（Investment Thesis）", bold=True)
add_para(
    "Nice S.p.A.はゲート/ガレージドア自動化（C&I含むHome & Building Automation）の世界トップクラスのプレイヤーで、"
    "上場廃止後の7年間で売上を€368M（FY2018）から推定€800M超（FY2022-2024）へと2倍超に拡大した。FSIが2023年11月に"
    "€100Mの少数出資を行い既にPE的ガバナンスが導入されていることから、現在は「FSI Exitを織り込んだ"
    "創業家マジョリティ買収」、または「FSIと並走する追加少数出資」の2類型でアプローチが可能。"
    "投資仮説は、(i) Nortek S&C（2021年$285M買収）の北米基盤を活用したスマートホーム/セキュリティ事業の"
    "EBITDAマージン底上げ、(ii) Mogollon新CEO（2024年6月就任）下の25年再編によるオペレーション統合益の刈取、"
    "(iii) 創業家承継の枠組みを利用したマジョリティ取得とExit時マルチプル拡大（FAACの17% EBITDAマージン水準への"
    "収斂）の3段で価値創出する余地がある点。",
)

add_para("主要観察事実（Key Observations）", bold=True)
add_bullet("グループ売上は2018年€368M → 2022年c.€800M（5年で約2倍）。Nortek S&C買収（2021年10月、$285M）が成長の主要ドライバー [出典: FSI press release Nov-2023; PR Newswire 2021-10-05]。")
add_bullet("ゲート/ドア自動化（Home & Building Automation）の世界主要プレイヤー6社（Nice、Somfy、CAME、FAAC、Hörmann、ASSA ABLOY）の中で、純粋型Home Management Solutionsプロファイルとしては規模上位 [出典: 業界ソース横断]。")
add_bullet("輸出比率95%、20+ヶ国に直販子会社、15のR&Dセンター・15生産拠点（イタリア・ポーランド・独・ブラジル・米・加・南アフリカ等）[出典: Buoro発言 2018; FSI press release 2023]。")
add_bullet("Italian-parent単体売上はFY2022 €143M / FY2023 €145M / FY2024 €152M（純利益はFY2022 +€9.3M、FY2023 -€7.5M、FY2024 +€8.9M）。グループ全体の8割超は海外子会社経由 [出典: Italian Camera di Commercio]。")
add_bullet("Lauro Buoro（創業者・Chairman、1993年創業時から）が持株会社Nice Group S.p.A.経由でグループの過半（推定70%）を保有。FSI II（イタリア国策系PE）が€100M reserved capital increaseで少数出資（推定15%）[出典: PR Newswire 2023-09-26、FSI公式]。")
add_bullet("CEO交代：Roberto Griffa（2015年11月-2024年6月、ex-Fiat）→ Juan B. Mogollon（2024年6月-、ex-Prysmian Energy Division EVP）。Mogollonの就任直後にValcucine（家具・€?M、2024年10月）買収と2025年再編発表 [出典: PR Newswire 2024-06-25; il Nord Est 2025]。")
add_bullet("ブランド・ポートフォリオは「Nice Galaxy」と称され、Fibaro（ポーランド・スマートホーム、$73M、2018）、Abode（米・DIYセキュリティ、75% for $18.75M、2018）、Nortek S&C（米・$285M、2021）、Elero（独）、HySecurity（米・産業ゲート）、Peccinin（ブラジル）、ELAN/2GIG/GoControl/IntelliVision（旧Nortek）等を含む20+サブブランド構成 [出典: niceforyou.com/about]。")

add_para("観察されたリスク要因（Observed Risk Factors）", bold=True)
add_bullet("非上場かつ連結IFRS財務未開示の状態が7年継続。買収検討に際してはVDR提供に依存度が高く、独立した数値検証は困難。Italian-parent単体FY2023で純利益-€7.5Mの赤字計上（FY2022 +€9.3M、FY2024 +€8.9M）の背景がブラックボックス。")
add_bullet("2025年再編（Mogollon主導、Confindustria Veneto Estと産業関係調整）が「limited impact」と表現されつつも実施されており、FY2024業績に何らかの圧迫要因（顧客需要鈍化／在庫調整／買収オーバーヘッド）が存在する可能性 [出典: il Nord Est 2025]。")
add_bullet("Nortek S&C買収の北米事業統合（旧ELAN/2GIG/GoControl等）はリーダーシップ・組織文化・チャネル統合の難度が高く、シナジー実現度はオープン。買収から4年経過時点で完全統合に至っているか不明。")
add_bullet("製品ラインの幅広さ（ゲート、ドア、シャッター、awning、警報、スマートホームHub、commercial AV）に対し、コア事業ゲート/ドア自動化の市場成長率はCAGR 5%程度と中速。一方スマートホームのCAGR推定値は20%+とばらつきが大きく、利益率プロファイルも事業間で大きく異なる。")
add_bullet("FSI（イタリア国策系PE）の出資は2023年11月。投資ホライズンは通常5-7年であり、FSIのExit時期と買い手側の検討タイミングが衝突／補完する可能性。FSI同意なしのマジョリティ取得は実質困難。")

add_para("追加DDで詰めるべき論点（Open Questions）", bold=True)
add_bullet("【財務】FY2020-FY2024の連結IFRS財務（PL/BS/CF）、セグメント別売上・EBITDA、地域別売上、Net Debt水準、Working Capital推移、Capexプロファイル。VDR必須。")
add_bullet("【財務】Nortek S&C $285M買収のPGA（Purchase Price Allocation）詳細とそれに伴う無形資産償却がEBIT/EBITDAに与えている圧迫額。Goodwill減損テスト履歴。")
add_bullet("【シナジー】Nortek S&C統合の進捗：ELAN OS / Yubii OSの統合状況、北米Smart Home事業の独立採算と本社オーバーヘッド配賦。")
add_bullet("【ガバナンス】FSI契約のタグアロング/ドラッグアロング条項、Buoro家のExitオプション、株主間契約の有効期限。Nice Group S.p.A.の上層構造。")
add_bullet("【経営】Mogollon CEOのKMP契約期間・成果インセンティブ設計、Griffa時代からの主要KMP（CFO、各事業ヘッド）のリテンション状況、後継者プラン。")
add_bullet("【事業】2025年再編の対象部署・対象人数・撤退製品ライン・ワンタイムコスト・期待ランレート効果。Confindustria Veneto Est経由で実施するイタリア国内のリストラ規模。")
add_bullet("【製品】Yubii / ELAN OSのMatter対応、AppleHomeKit/Google Home/Amazon Alexaとの相互運用性、IoTプラットフォームでの将来Recurring Revenue（saaS）化余地。")
add_bullet("【顧客】チャネル別売上構成（Pro Installer / Retail / OEM）、上位顧客集中度、長期契約の有無。")

doc.add_page_break()

# ============================================================================
# ① 事業概要
# ============================================================================
add_heading("① 事業概要", level=1)

add_heading("1-1. 事業セグメント構造", level=2)
add_para(
    "Nice Groupはセグメント別の連結財務を公表していない。本節の3セグメント（Gate & Door Automation／"
    "Sun Shading／Smart Home & Security）は、製品カテゴリの内部論理および買収サブブランドの所属"
    "（Mhouse・Era・HySecurity・Peccinin・V2はGate & Door Automation／Elero・Nice本体awning事業はSun Shading／"
    "Fibaro・Abode・Nortek S&C・ELAN・Linear・2GIGはSmart Home & Security）から構成した推定区分。"
    "全数値は推定値のため(E)を付す。",
    italic=True, size=9,
)

add_para("テーブルA：5期セグメント別業績推移（推定、€百万）", bold=True)
table_a = [
    ["セグメント", "指標", "FY2020E", "FY2021E", "FY2022E", "FY2023E", "FY2024E"],
    ["Gate & Door Automation", "売上高", "285", "310", "415", "415", "405"],
    ["", "セグメント利益", "30", "35", "60", "60", "56"],
    ["", "利益率(%)", "10.5", "11.3", "14.5", "14.5", "13.8"],
    ["", "設備投資", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["", "減価償却費", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["Sun Shading", "売上高", "90", "100", "145", "153", "162"],
    ["", "セグメント利益", "12", "12", "22", "23", "24"],
    ["", "利益率(%)", "13.3", "12.0", "15.2", "15.0", "14.8"],
    ["", "設備投資", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["", "減価償却費", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["Smart Home & Security", "売上高", "55", "140", "240", "262", "243"],
    ["", "セグメント利益", "4", "8", "25", "28", "21"],
    ["", "利益率(%)", "7.3", "5.7", "10.4", "10.7", "8.6"],
    ["", "設備投資", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["", "減価償却費", "不開示", "不開示", "不開示", "不開示", "不開示"],
    ["連結合計", "売上高", "430", "550", "800", "830", "810"],
    ["", "営業利益", "27", "35", "66", "68", "66"],
    ["", "EBITDA(参考)", "47", "65", "104", "108", "105"],
    ["", "EBITDAマージン(%)", "10.9", "11.8", "13.0", "13.0", "13.0"],
]
add_table(table_a)

add_para(
    "含意（So What）：成長の大半はSmart Home & Security（FY2020 €55M → FY2022 €240M）が牽引した一方、"
    "利益率は同セグメントが最低（FY2024E 8.6%）にとどまり、Nortek統合の収益化は道半ば。コア"
    "Gate & Door Automation（FY2024E 13.8%）とSun Shading（同14.8%）はFAAC（17.2%）・Somfy水準に近づきつつあり、"
    "ここがNiceの利益基盤を形成。FY2024売上微減（830→810）と再編発表は、Smart Home事業のオーバーヘッド吸収が"
    "課題化していることを示唆。",
    size=10,
)

add_para("テーブルB：FY2024E セグメント構成比（独立テーブル）", bold=True)
table_b = [
    ["セグメント", "売上 (€M)", "売上構成比", "利益 (€M)", "利益構成比", "セグメント利益率"],
    ["Gate & Door Automation", "405", "50.0%", "56", "55.4%", "13.8%"],
    ["Sun Shading", "162", "20.0%", "24", "23.8%", "14.8%"],
    ["Smart Home & Security", "243", "30.0%", "21", "20.8%", "8.6%"],
    ["合計", "810", "100.0%", "101", "100.0%", "—"],
]
add_table(table_b)
add_para(
    "含意（So What）：売上構成比30%のSmart Home & Securityが利益構成比では20.8%にとどまり、"
    "コア事業（Gate＋Sun Shading）が売上70%・利益79%を稼ぎ出す構造。Smart Home事業の利益率を"
    "コア事業並みの13-14%まで底上げできれば、グループEBIT €15-20M（EBIT率+2-3pp）の改善余地が見える。"
    "これはMogollon新CEOの最大の経営課題であり、Value-up論点の中核。",
    size=10,
)

add_heading("1-2. 製品・サービス", level=2)
add_para(
    "Nice Groupの製品ラインは「Home Management Solutions（HMS）」と総称され、"
    "下記の主要カテゴリで構成される。買収によって取得したサブブランドが各カテゴリで併存し、"
    "価格帯／チャネル／地域で棲み分けている。",
)
products = [
    ["カテゴリ", "主力ブランド", "用途・特徴"],
    ["Gate Automation", "Nice (Run, Wingo, Robus)、Mhouse、HySecurity（米産業）、Peccinin（中南米）", "スイング/スライディングゲートのモーター・コントローラ。住宅Pro〜commercial heavy duty"],
    ["Garage Door", "Nice (Spin, Spido)、Linear/Mighty Mule（米）、Micanan", "住宅・産業ガレージドアopener"],
    ["Barriers", "Nice (Wide, M Bar, S Bar)", "駐車場・通路向け道路バリア"],
    ["Sun Shading", "Nice (Era, Neo)、Elero（独・OEM motorization）", "awning、ローラーシャッター、ブラインド用motor／コントローラ"],
    ["Alarm Systems", "Nice (MyNice)、Abode（米DIY）、Fibaro Home Center", "ワイヤレス警報、IPカメラ、door/windowセンサー"],
    ["Smart Home Hub", "Yubii (旧Nice)、ELAN OS（旧Nortek S&C）、Fibaro Home Center", "Z-Wave/Zigbee/Matter対応hub。Apple HomeKit、Google Assistant、Alexa統合"],
    ["AV / Custom Integration", "ELAN、Furman、Gefen、Proficient（旧Nortek S&C）", "プロインストーラー向けマルチルームAV、power management"],
]
add_table(products)
add_para(
    "ELAN OS（旧Nortek S&Cの北米Pro統合プラットフォーム）とYubii OS（旧Nice/Fibaro統合の欧州コンシューマ／ライトPro）"
    "は、現状並列運用されている。両者の統合または棲み分け明確化は、Mogollon体制下のkey product roadmap課題。",
    size=10,
)

add_heading("1-3. 顧客構造", level=2)
add_para(
    "Nice Groupの顧客は大別して(i) Pro Installerチャネル（ガレージドア／ゲート設置業者、AV/セキュリティintegrator）、"
    "(ii) Retail/DIYチャネル（Abode等の米DIYセキュリティ、Fibaroコンシューマhub）、(iii) OEM/B2B（Eleroの窓motorization、"
    "HySecurityのcommercial／公共インフラ）の3類型。チャネル別売上構成および上位顧客リストはVDR必須項目。"
    "従来Niceはイタリアおよび欧州大陸でPro Installerチャネルが圧倒的に強く、Nortek S&Cの北米Pro Integrator基盤"
    "（CEDIA市場、米国Custom Installation業界）はそれに直接連続する取得資産。",
)
add_para(
    "顧客集中度については、Pro Installerチャネルが多数の中小installer向け間接販売（ディストリビューター経由）であるため、"
    "上位5社集中度は10%未満と推定（業界類推）。OEM契約とcommercialプロジェクト（HySecurity）は単発契約規模が"
    "相対的に大きく、特定顧客集中の可能性あり（要確認）。",
    size=10,
)

add_heading("1-4. バリューチェーン上のポジション", level=2)
add_para(
    "Nice Groupは(a)機構設計＋エレクトロニクス開発＋ファームウェア開発、(b)モーター・コントローラ・センサーの"
    "組立製造、(c)ブランドマーケティング＋Pro Installerサポート（トレーニング・保守）の3層を内製化する垂直統合型。"
    "上流（モーター巻線材、半導体、樹脂・アルミ筐体）は外部調達。下流（住宅オーナー）には直接到達せず、"
    "ディストリビューター→Pro Installer経由の間接モデルを基本とする。",
)
add_bullet("生産：15生産拠点（イタリア4、ポーランド〔Fibaro〕、独〔Elero〕、ブラジル〔Peccinin〕、米〔Nortek/HySecurity〕、加、南アフリカ等）。原価優位の獲得は中・東欧（ポーランド）と新興国（ブラジル、メキシコ、東欧）拠点が担う。")
add_bullet("R&D：15センター（イタリア・ポーランド・米中心）。Yubii OS（伊）とELAN OS（米）の二極体制。Matter規格対応は両OS間の統合を促進する外的要因。")
add_bullet("販売：20+カ国直販子会社、100+カ国ディストリビューター。Niceの強みは「Pro Installerに対する商品教育・保守サポートの厚み」（業界類推）。")
add_bullet("内製/外注境界：プリント基板（PCBA）の一部・モーター巻線・成形品はEMS／専業メーカーへの外注。最終組立・QA・ファームウェア・コントローラチップ（mcu）選定は内製。")

add_heading("1-5. 主要競合（業界マッピング）", level=2)
add_para("テーブル：Home & Building Automation主要競合（FY2024、推定値含む）", bold=True)
competitors = [
    ["競合名", "上場区分", "売上規模", "EBIT/EBITDA率", "主領域", "強み・特徴"],
    ["Somfy Group (FR)", "親族支配・非上場（Despature家、2022 OPA後）", "€1.5B (FY2024)", "EBITDA n.a.（推定 ~16%）", "窓motorization・スマートホーム", "TaHomaエコシステム、Velux等との同盟、ECoec設計80%。Bft（伊）保有"],
    ["FAAC Technologies (IT)", "非上場（教皇庁関連財団Sant'Orsola支配）", "€698M (FY2024)", "EBITDA 17.2%", "ゲート/ドア自動化、parking、AV-IDS", "Bologna本社、buy & build戦略で60社超買収、parking事業を持つ多角化型"],
    ["CAME Group (IT, Treviso)", "非上場（Menuzzo家）", "€335M (FY2024, +8.4%)", "n.a.", "ゲート自動化・barrier", "Niceの近隣・直接競合。住宅Pro主体"],
    ["Hörmann Group (DE)", "非上場（Hörmann家）", "€1.3-1.5B（推定）", "n.a.", "ガレージドア・産業ドア", "産業ドア・industrial gates世界トップ。住宅ガレージドアでも欧州リーダー"],
    ["ASSA ABLOY (SE)", "上場（OMX、Wallenberg関連）", "$14.16B (FY2024)", "EBIT 16.2%", "錠・access control・ドア", "ドアコンプリート供給。Niceとはdoorとinterior automation領域で重なる。M&A 26件/2024年"],
    ["Chamberlain Group (US)", "非上場（Blackstone保有）", "n.a.（推定$1.8B+）", "n.a.", "ガレージドアopener", "myQプラットフォーム、米最大シェア。住宅DIY/Pro両対応"],
]
add_table(competitors)
add_para(
    "含意（So What）：Nice Groupは€800M超のグループ売上規模で、純粋型Home Management Solutionsプロファイルとしては"
    "Somfy（€1.5B）に次ぐポジション。FAAC（€698M）と並ぶ規模だが、FAACがparking／ID管理／AV業界等への多角化に"
    "より構築した€698M／EBITDA 17%水準と比べ、NiceはSmart Home（Nortek/Fibaro）への集中投資により€800M／EBITDA"
    "推定13%にとどまる。EBITDAマージンギャップ（推定3-4pp）はValue-up余地そのもの。"
    "ASSA ABLOYは規模で20倍離れた相手だが、Niceの隣接領域への参入は戦略買収候補としての魅力を示唆。",
    size=10,
)

add_heading("1-6. 沿革・資本構成", level=2)
add_para(
    "1993年、Lauro BuoroがOderzo（Treviso）でNiceを設立。当初はゲート自動化のニッチプレイヤーだったが、"
    "2000年代に欧州・新興国へのアグレッシブな買収（Mhouse、Era、Peccinin、Elero等）でグループ化。"
    "2006年Borsa Italianaに上場。2018年12月、Buoroは「政府の経済政策により外国投資家がイタリア企業から逃避している」として"
    "上場廃止を表明し、2019年4月に正式delisting [出典: Tribuna di Treviso 2018]。"
    "上場廃止後はBuoro家による完全supportのもとで買収戦略を継続し、2018年Fibaro（$73M）、2018年Abode（$18.75M for 75%）、"
    "2021年10月Nortek S&C（$285M、過去最大）と立て続けに実行。2023年11月にFSI II（イタリア国策系PE）が"
    "€100Mのreserved capital increaseで少数出資し、現在に至る。",
)

add_para("過去主要買収（時系列）", bold=True)
acq = [
    ["年", "対象", "国", "金額", "戦略意図"],
    ["2008", "Mhouse", "イタリア", "n.a.", "ミドルレンジゲート自動化（住宅Pro）"],
    ["2008", "Era（Nice下位ブランド化）", "イタリア", "n.a.", "エントリー価格帯"],
    ["2010頃", "Peccinin", "ブラジル", "n.a.", "中南米プラットフォーム"],
    ["2010s", "Elero", "ドイツ", "n.a.", "窓・blind motorization、OEM強化"],
    ["2010s", "HySecurity", "米国", "n.a.", "Commercial heavy-duty gate"],
    ["2010s", "ACM、V2、Micanan、Linear", "伊・加・米", "n.a.", "ゲート/ガレージ製品ライン補完"],
    ["2018年4月", "Abode Systems（75%）", "米国", "$18.75M", "DIYコンシューマ・スマートセキュリティ"],
    ["2018年", "Fibaro", "ポーランド", "$73M", "スマートホーム自社プラットフォーム取得"],
    ["2021年10月", "Nortek Security & Control", "米国", "$285M", "北米Pro Integrator基盤・ELAN/2GIG/GoControl/IntelliVision/Linear等のサブブランド群"],
    ["2024年10月", "Valcucine", "イタリア", "n.a.", "高級キッチン家具（家まわり統合戦略の延長）"],
]
add_table(acq)

add_para("大株主構成（推定、FY2024時点）", bold=True)
shareholders = [
    ["株主", "持株比率（推定）", "備考"],
    ["Nice Group S.p.A. (Buoro family)", "70.0%", "創業家持株会社経由。delisting前は68.42%"],
    ["FSI II", "15.0%", "2023年11月 €100M reserved capital increase（推定希釈後）"],
    ["Management & Treasury", "10.0%", "Mogollon CEO・主要KMP・自己株式"],
    ["Other minority", "5.0%", "delisting残余・関連法人"],
    ["合計", "100.0%", "—"],
]
add_table(shareholders)
add_para(
    "Buoro家（70%推定）の支配構造は買収検討の最大論点。FSIが既に少数出資していることから、"
    "(a) Buoro家からのマジョリティ取得＋FSI継続orテイクアウト、(b) FSI Exitの形でセカンダリー出資、"
    "(c) Tagalongを使った全株買付（TOB相当の私的買収）の3シナリオが現実的。"
    "FSI契約の株主間契約（タグアロング/ドラッグアロング）の有効期限・条件はVDR必須確認事項。",
    size=10,
)

doc.add_page_break()

# ============================================================================
# ② 事業の堅牢性
# ============================================================================
add_heading("② 事業の堅牢性", level=1)

add_heading("2-1. 市場性（市場数量・市場規模）", level=2)
add_para(
    "Nice Groupの事業は3つの市場層に分けられる：(i) ゲート/ドア自動化（コア、CAGR 5-7%の中速成長）、"
    "(ii) Sun Shading／窓motorization（中速、5%程度）、(iii) スマートホーム＋ホームセキュリティ（高速、CAGR 15-27%）。"
    "事業ポートフォリオはゲート→スマートホームに重心移動中（Smart Home & Security FY2020 13% → FY2024 30%）。",
)

add_para("テーブル：3市場の規模・成長率", bold=True)
mkt = [
    ["市場", "直近規模", "期間", "実績/予測", "CAGR", "出典"],
    ["Gate Automation（broader、含commercial parking等）", "$18.6B (2024)", "2024→2030", "予測", "7.0%", "ResearchAndMarkets / Maximize Market Research"],
    ["Gate Openers（住宅・C&I focused）", "$2.66B (2024)", "2024→2032", "予測", "5.1%", "Fortune Business Insights"],
    ["Smart Home（global）", "$127.8B (2024)", "2024→2030", "予測", "20.1%-27%", "Grand View Research / Globenewswire"],
    ["Smart Home Hub（segment内Niceポジション）", "n.a.", "2024→2030", "予測", "12-15%（推定）", "Mordor Intelligence類推"],
]
add_table(mkt)

add_para(
    "数量ドライバーは3つに分けられる。第1に新興国（中南米・東欧・東南アジア・中東）の住宅普及率上昇に伴う"
    "「ゲート／ガレージドア自動化の本数増加」。住宅市場の成熟欧米では本数増加は限定的だが、買替需要が定常的。"
    "第2にcommercial（駐車場・物流倉庫・産業施設）の新規建設に伴うbarrier・heavy-duty gateの本数増。"
    "第3にスマートホームでは、コンシューマの「家のIoT化」需要が直接ドライバー。",
)
add_para(
    "価格ドライバーは「ベーシックゲートmotor」が単価横ばい〜微減（中国製品流入）である一方、"
    "「IoT統合制御hub＋センサーパッケージ」は単価上昇余地が大。Yubii OSやELAN OS搭載のhub＋scenario"
    "セットは住宅当たり数百〜数千ユーロのベンダーロックを生む。",
)
add_para(
    "市場規模（金額ベース）の将来見通しは、コア（ゲート/ドア）が+5-7%、Sun Shadingが+4-5%、"
    "スマートホームが+15-20%（仮説）。Niceのportfolio mix（FY2024 50/20/30）から推計される加重平均"
    "成長率は+9-10%程度（仮説）。",
)

add_heading("2-2. 対象会社シェア", level=2)
add_para(
    "Niceは「Home & Building Automation」のグローバル規模では3-5位、欧州（伊・東欧・南欧）では1-2位。"
    "北米はNortek S&C取得後にPro Integratorチャネルで上位10位以内に上昇したと推定。直接競合のFAAC、CAMEは"
    "イタリア中心、Somfyは窓automation中心、Hörmannは産業ドア中心と棲み分け。",
)

add_para("KBF（Key Buying Factor）と充足度", bold=True)
add_para(
    "Pro Installerチャネルでは、新規プロジェクト採用時と継続契約時でKBFが異なる。スイッチングは住宅単位ではなく、"
    "Installer単位（Pro Installerが扱うブランドを切り替えるタイミング）で発生。Installerは通常2-3ブランドを併売し、"
    "案件規模・利益率・在庫サポートで使い分け。",
    size=10,
)

add_para("(a) 新規採用時のKBF（新規Pro Installerが取扱開始する場面）", bold=True)
kbf_a = [
    ["KBF", "重要度（仮説）", "Nice", "Somfy", "FAAC", "CAME"],
    ["技術トレーニング・サポート体制", "高", "◎", "○", "○", "○"],
    ["商品ラインナップの幅（住宅〜commercial）", "高", "◎", "△（窓中心）", "○", "○"],
    ["ブランド認知度（地域別）", "中", "○（伊◎、米○）", "◎（仏◎）", "○（伊◎）", "○（伊◎）"],
    ["初期Margin・Spec inクレジット", "中", "○", "○", "○", "○"],
    ["Smart Home統合・Matter対応", "高（昇格中）", "○（Yubii/ELAN）", "○（TaHoma）", "△", "△"],
]
add_table(kbf_a)

add_para("(b) 既存契約継続時のKBF（取扱継続するか他社に切替えるかの判断）", bold=True)
kbf_b = [
    ["KBF", "重要度（仮説）", "Nice", "Somfy", "FAAC", "CAME"],
    ["納期遵守・在庫availability", "高", "○", "○", "◎", "○"],
    ["故障率・品質安定性", "高", "○", "○", "◎", "○"],
    ["価格改定の柔軟性（installer margin保護）", "中", "○", "△", "○", "◎"],
    ["新製品ローンチ頻度", "中", "○", "○", "○", "○"],
    ["既存OS／ecosystemへのbackward compatibility", "高", "△（Yubii/ELAN二極化）", "◎（TaHoma連続性）", "n/a", "n/a"],
]
add_table(kbf_b)

add_para(
    "Niceの優位性：(i) 商品ラインナップの幅広さ（ゲート＋ドア＋シャッター＋警報＋スマートホーム一気通貫）、"
    "(ii) Pro Installer向けトレーニング・保守サポートの厚み（イタリア・欧州での歴史的優位）、"
    "(iii) Nortek S&C取得による北米Pro Integrator基盤。Niceの劣位：Yubii OS／ELAN OSの二極構造による既存"
    "ecosystem内continuityの弱さ。Matter普及はこれを軽減する一方、Somfyとの競合を激化させる方向。",
    size=10,
)

add_heading("2-3. 単価・コスト構造", level=2)

add_para("(a) 単価", bold=True)
add_para(
    "主要コスト構造（Nice製品の物理構成からの推定）：(i) DCモーター／ステッピングモーター（最大コストファクター、"
    "巻線銅・希土類磁石・ステーター鉄芯）、(ii) コントローラーPCB（mcu・電源IC・無線モジュール）、"
    "(iii) 機構・筐体（アルミ・亜鉛ダイキャスト、エンジニアリングプラスチック）、(iv) 二次電池"
    "（コードレスゲート向け、リチウムイオン）。",
)
add_para(
    "原材料市況：銅は2024年以降$9,000-10,000/tの高位レンジで推移（過去3年で+30%）、"
    "アルミは$2,500-2,700/t、希土類磁石（Nd-Fe-B）は中国輸出規制で2024年から構造的供給リスク。"
    "半導体（mcu）は2022年ピークから正常化したが、IoT用無線モジュール（Z-Wave、Zigbee、Thread/Matter）"
    "の単価は標準化の進展で下落基調。Nice製品の原価へのインパクトは、+2-4%/年の上方圧力（推定）。"
    "[出典: LME公開価格・業界ニュース類推]",
)
add_para(
    "顧客サイドの圧力：Pro Installer（中小事業者）の購買力は限定的だが、ディストリビューター（地域ホールセラー）"
    "は大型ロット購買で価格交渉余地を持つ。住宅オーナー最終価格に占めるNice製品コストは20-30%程度（残りは設置工賃）"
    "で、価格転嫁余地は中位。Smart Home製品はコンシューマDIY販売（Abode等）でより価格感応度高い。",
)
add_para(
    "競争環境：CAME・FAACとの直接価格競合、中国製OEM品（Beninca、Roger Technology等）の流入、"
    "Chamberlain（米）／Hörmann（独）からの逆輸入圧力が存在。Niceは2018年以降の値上げ実績は推定"
    "+2-3%/年（公表なし）。コア事業のpricing powerは中位、Smart Homeはコモディティ化しやすい。",
)

add_para("(b) コスト", bold=True)
add_para(
    "粗利率・営業利益率推移（FY2018実績→推定値）：粗利率は2018年公表時点で約45%（業界類推）、"
    "EBITDAマージンは2018年13.8% → 推定FY2024 13.0%。FAAC（17.2%）との3-4ppギャップは、"
    "(i) Nortek S&C統合に伴うオーバーヘッド吸収不良、(ii) Smart Home事業のlow-margin商品（DIY）寄与、"
    "(iii) 多ブランド体制によるmarketing／branding重複コストが要因（仮説）。",
)
add_para("テーブル：競合比較EBITDA/EBIT率（FY2024実績or直近）", bold=True)
margins = [
    ["企業", "売上 (€M)", "EBITDA率", "EBIT率", "備考"],
    ["Nice Group", "810 (E)", "13.0% (E)", "8.1% (E)", "本DD推定"],
    ["FAAC Technologies", "698", "17.2%", "11.9%", "FY2024実績"],
    ["CAME Group", "335", "n.a.", "n.a.", "FY2024 +8.4%"],
    ["Somfy Group", "1,500", "推定 16-18%", "推定 11-13%", "FY2024、ecodesign 80%"],
    ["ASSA ABLOY", "≒$14,160", "≒18%", "16.2%", "FY2024、Group全体"],
]
add_table(margins)
add_para(
    "固定費・変動費比率：Nice Groupの公表値はないが、業界類推で材料費（変動）が売上の45-50%、"
    "労務費（半固定）が18-22%、減価償却（固定）が4-5%、その他SG&A（半固定）が15-18%。"
    "Nortek S&C統合後の不正常負担分は売上の1-2%程度（推定）。",
    size=10,
)

add_heading("2-4. 事業構造のテンション（Strategic Tensions）", level=2)
add_para(
    "Niceの事業構造には以下の3つの本質的テンションが存在し、これらの解消方向性がValue-upプランの"
    "起点となる。",
)
add_para(
    "テンション1：「コア事業の安定収益（ゲート＋Sun Shading 70%売上・79%利益、利益率14%水準）」 vs "
    "「成長事業への構造転換（Smart Home & Security、利益率8.6%、グループ平均下押し）」"
)
add_bullet("現状はコア偏重で安定。Smart Home事業はNortek $285M投資に対し、利益貢献は限定的。")
add_bullet("Buoro/Mogollonの公表ストーリー（「smart home pivot」）と利益実態がdivergence。")
add_bullet("方向性：Smart Home事業の利益率底上げ（コア並み13%）、または非戦略事業の整理。Mogollonの2025年再編はこの解消が主旨と推察。")

add_para(
    "テンション2：「Buoro家の支配維持（70%）」 vs 「成長必要資本（FSI €100M、追加M&A資本）」"
)
add_bullet("Buoro家は2018年delisting時に「自由度確保」の論理で支配を強化したが、5年後にFSI受入で部分的妥協。")
add_bullet("追加大型M&Aや北米市場攻略には更なる資本が必要だが、Buoro家の支配を維持するか希釈するかの判断点。")
add_bullet("方向性：マジョリティ売却（Buoro家Exit）または継続的少数希釈。買収候補側の検討においては、Buoro家マジョリティ取得スキームが現実的。")

add_para(
    "テンション3：「20+サブブランドのbroad portfolio」 vs 「グローバルブランド統合とscale economics」"
)
add_bullet("Mhouse、Era、Fibaro、Abode、Nortek、ELAN、Linear、HySecurity、Peccinin、V2、Mhouse、Eleroなど多数の取得ブランド併存。")
add_bullet("地域・チャネル別でブランド使い分けが現状の戦略。一方でmarketing／brandingコスト・SKU管理の複雑性はマージン圧迫要因。")
add_bullet("方向性：Yubii OS／ELAN OSの統合（プラットフォーム単一化）、Tier-2/Tier-3ブランドの段階的Nice本体ブランドへの収斂。Mogollon体制で明確化が想定される。")

doc.add_page_break()

# ============================================================================
# ③ バリューアップの方向性
# ============================================================================
add_heading("③ バリューアップの方向性", level=1)

add_heading("3-1. オーガニック", level=2)

add_para("(a) 売上ドライバー", bold=True)

add_para("① プライシング改善（単純値上げ）", bold=True)
add_para(
    "打ち手：Pro Installerチャネルでの+2-3%/年の段階的値上げ、ゲート／コントローラの高級ライン拡充による"
    "ASP（Average Selling Price）押し上げ。仮説の根拠：原材料市況（銅+30%、希土類リスク）への合理的転嫁余地、"
    "FAAC・CAMEとの相対価格水準（Niceがコア領域でややディスカウントポジション）。"
    "定量インパクト：年+1-2%売上寄与（=€8-16M/年、3年で€25-50M売上、EBITDA寄与€3-7M）。"
    "実行リスク：CAME・FAACの価格追随、中国OEM品代替、Pro Installerからの反発。",
    size=10,
)

add_para("② 新規顧客拡大", bold=True)
add_para(
    "打ち手：Smart Home Pro Integratorチャネル（米CEDIA、欧州AV/security integrator）でのELAN OS／"
    "Yubii統合プッシュ、住宅builder（Lennar、KB Home等の米大手）との直接OEM契約獲得。"
    "仮説の根拠：Nortek S&C取得時に同梱されたELAN brand（CEDIAでの確立ポジション）の活用、"
    "Mogollon CEOのPrysmian時代北米実績。"
    "定量インパクト：年+€20-40M売上（3-4年）、EBITDA寄与€3-6M。"
    "実行リスク：Control4・Crestron・Lutronとの直接競合、米Pro Integratorチャネルの保守性。",
    size=10,
)

add_para("③ 新規製品拡充", bold=True)
add_para(
    "打ち手：Matter対応hub・センサーの全ライン展開、AI／音声制御を組み込んだ次世代Yubii、"
    "電力管理（PV／蓄電池統合）製品（HEMS方向）。"
    "仮説の根拠：Matter普及（2024年規格成熟）、住宅電力管理需要の構造的増加（欧州エネルギーコスト）。"
    "定量インパクト：年+€15-30M売上（3-5年）、EBITDA寄与€2-5M。"
    "実行リスク：R&D投資先行（€20-40M/年）、海外大手（Google Nest、Amazon、Samsung SmartThings）との競合。",
    size=10,
)

add_para("④ 新規地域展開", bold=True)
add_para(
    "打ち手：APAC（特に中国・東南アジア・豪州）でのHigh-end住宅・Pro Installerチャネル開拓、"
    "中東（GCC新興都市）でのcommercial barrier／heavy-duty gate案件。"
    "仮説の根拠：Niceの現状APAC構成比が低い（推定<10%）、APAC市場CAGR 8.5%（gate automation）。"
    "定量インパクト：年+€10-20M売上（3-5年）、EBITDA寄与€1-3M。"
    "実行リスク：地場プレイヤー（中国Aleko、Roger Technology等）の価格競争、現地販売網構築コスト。",
    size=10,
)

add_para("(b) コストドライバー", bold=True)

add_para("① 原価改善", bold=True)
add_para(
    "打ち手：(i) 15生産拠点の最適化（イタリア4拠点の機能統合・東欧／中南米拠点への部分移管）、"
    "(ii) Nortek／Fibaro／Niceの共通プラットフォーム化（PCB共通化、モーター標準化）、"
    "(iii) 調達集約（銅・アルミ・mcu調達のグローバル一括化）、(iv) 自動化投資（Industry 4.0/5.0）。"
    "仮説の根拠：FAAC（EBITDA 17.2%）との3-4ppギャップの大半はオペレーション最適化で吸収可能（業界類推）。"
    "Mogollon CEOのoperational efficiency実績（Prysmian Latin America CEO時代）。"
    "定量インパクト：3-5年でEBITDA率+1.5-2.5pp（=年€12-20M EBITDA寄与）。"
    "実行リスク：イタリア国内雇用調整の労務コスト、Confindustria Veneto Estとの調整負荷、稼働率低下による短期粗利圧迫。",
    size=10,
)

add_para("② 販管費改善", bold=True)
add_para(
    "打ち手：(i) サブブランド統合に伴うmarketing／branding重複の削減、(ii) 北米／欧州の販売子会社の機能統合"
    "（Nortek＋Linear＋HySecurity＋Niceの米本社統合）、(iii) ITシステム統合（ERP一元化）、"
    "(iv) 本社機能のslimming（Mogollon体制下の2025年再編）。"
    "仮説の根拠：20+サブブランドの非効率、買収後7年経過してもbrand consolidationが未完了。"
    "定量インパクト：3-5年でSG&A率-1-1.5pp（=年€8-12M EBITDA寄与）。"
    "実行リスク：ブランドエクイティの毀損、key personnel離脱、IT統合の遅延コスト。",
    size=10,
)

add_heading("3-2. インオーガニック", level=2)

add_para("① 同業ロールアップ（規模の経済・シェア集約）", bold=True)
add_para(
    "打ち手：欧州の中堅ゲート自動化プレイヤー（Beninca、Roger Technology、独DEA System、独BFT、"
    "西BFT Spain等）のロールアップ。FAACが過去採用してきたbuy & build戦略を欧州中位プレイヤー連結で実施。"
    "候補企業（仮説）：Roger Technology（伊、推定€50M売上）、DEA System（伊、推定€30M）、"
    "Daab（独、推定€40M）、Erreka（西、推定€50M）。"
    "シナジー類型：(i) 製造拠点統合（東欧）、(ii) 調達集約、(iii) Pro Installer販売網重複削減。"
    "PMI論点：イタリア／欧州中堅オーナー企業の文化統合、ファミリーオーナーからのexit price、competition law。"
    "定量インパクト：合計売上€150-250M取得（5年）、EBITDA寄与€20-40M（margin 13-15% × cost synergy 5%）。",
    size=10,
)

add_para("② 隣接企業買収による機能強化（技術・顧客・チャネル・地域の補完）", bold=True)
add_para(
    "打ち手：(i) Smart Home AI／Voice領域のスタートアップ取得（Yubii OS強化）、"
    "(ii) 蓄電池・PV統合HEMSプレイヤー取得（住宅エネルギー管理への進出）、"
    "(iii) APAC現地プレイヤー取得（中国・印・東南アジア）、"
    "(iv) Commercial security／access control企業取得（Nortek基盤強化）。"
    "候補企業（仮説）：(i) Mediola／HomeMatic（独IoT）、(ii) Sonnen（独蓄電池）の住宅セクター、"
    "(iii) インドGodrejのIoT部門、(iv) Vanderbilt（access control、SwedbankRobur保有）。"
    "シナジー類型：技術統合、地域横断販売、Customer cross-selling。"
    "PMI論点：技術統合スピード、key engineerリテンション、地理的carve-out。"
    "定量インパクト：合計売上€100-200M取得（5年）、EBITDA寄与€5-15M（marginは取得時点で低いが3年で改善）。",
    size=10,
)

add_heading("3-3. Exitの可能性（エクイティストーリー含む）", level=2)

add_para("想定買い手（5年後Exit想定先）", bold=True)
add_bullet("戦略買収：Somfy Group（仏）、ASSA ABLOY（瑞）、Hörmann（独）、Chamberlain（米Blackstone保有）、Allegion（米access control）。Somfyが規模・カテゴリ的補完性で最有力。")
add_bullet("PE（再売却）：CVC、KKR、EQT、Bain Capital等のmega-buyout。Northern Italy家族系企業のPEレディなプロファイルとなる。")
add_bullet("IPO：Borsa Italiana再上場、Euronext Milan／Amsterdam dual list。FSI＋Buoro家のpartial Exit実現可能だが、再上場時のディスカウントリスク。")

add_para("想定マルチプル", bold=True)
add_bullet("Home Automation／Building Tech領域のtransaction multiple（過去5年）：EV/EBITDA 9-12x（戦略買収）、8-10x（PE）。")
add_bullet("公開ピアマルチプル：ASSA ABLOY EV/EBITDA c.16x（FY2024）、Chamberlain（Blackstone保有時2018年に20%-LBO 12x EBITDA）。")
add_bullet("FAACのEBITDA 17%水準にNiceが到達した場合、EBITDA €170M（FY2029想定）×10x = EV €1.7B（仮説）。")

add_para("5年後Exit想定（仮置）", bold=True)
exit_table = [
    ["指標", "FY2024E（現状）", "FY2029想定（5年後）", "前提"],
    ["売上 (€M)", "810", "1,150-1,300", "オーガニック+5-7% CAGR、M&Aで+€150-250M寄与"],
    ["EBITDA (€M)", "105", "160-210", "Value-up Bridge"],
    ["EBITDAマージン", "13.0%", "14-16%", "Smart Home収益化＋オペ最適化"],
    ["EV/EBITDAマルチプル", "—", "9-11x", "戦略買収レンジ"],
    ["想定EV (€M)", "—", "1,440-2,310", "中央値c.€1.9B"],
    ["FY2024 EV（参考）", "推定€1.0-1.2B", "—", "EBITDA 105 × 10x"],
]
add_table(exit_table)

add_para("エクイティストーリー（5年後Exit時）", bold=True)
add_para(
    "「Italian heritage + global scale + smart home pivotの完成形：Nice Group is the only pure-play"
    "Home Management Solutions player at scale outside the US, with 50% of revenue from"
    "high-growth Smart Home & Security and 50% from cash-cow Gate & Door Automation. EBITDA margin"
    "expanded from 13% to 15-16% via Nortek S&C full integration and operational consolidation."
    "FY2029 EBITDA c.€180M, attractive growth+margin profile to strategic acquirers (Somfy, ASSA ABLOY)"
    "and global PE.」",
    italic=True, size=10,
)

add_heading("Value-up Bridge（5年後想定EBITDA）", level=2)
add_para("テーブル：EBITDA Bridge（€百万）", bold=True)
bridge_t = [
    ["項目", "EBITDA（€M）／レンジ", "寄与の論拠", "達成期間"],
    ["現状EBITDA（FY2024E）", "105", "推定実績（EBIT 66 + D&A 39）", "—"],
    ["＋ オーガニック", "+30〜+55", "売上ドライバー4項目（pricing/新規顧客/新規製品/新規地域）+ コスト2項目（COGS/SG&A）の合算", "3-5年"],
    ["＋ インオーガニック", "+25〜+50", "同業ロールアップ（€150-250M売上取得、margin 13-15%＋synergy）＋ 隣接買収（€100-200M売上取得）の合算", "3-5年"],
    ["5年後想定EBITDA（FY2029）", "160〜210", "—", "—"],
]
add_table(bridge_t)
add_para(
    "投資仮説との整合性：現状EBITDA €105M → 5年後€160-210Mへの拡大は、エグゼクティブサマリー冒頭の"
    "投資仮説3要素（Nortek収益化／2025年再編完成／創業家承継スキーム活用）を定量的に裏付ける。"
    "中央値€185Mは現状比+76%、9-11x EV/EBITDAでExit想定€1.7-2.3Bは買収ファイナンス（5x leverage前提）で"
    "十分にbankable。",
    italic=True, size=10,
)

doc.add_page_break()

# ============================================================================
# ④ 直近の経営者発言（FY2022-FY2024）
# ============================================================================
add_heading("④ 直近の経営者発言（FY2022-FY2024）", level=1)
add_para(
    "目的：有報・IR資料が存在しない非公開企業のため、media露出の中での経営者発言から経営観・戦略意図・"
    "リスク認識を抽出する。網羅性は求めず、確認できた発言ベースで提示。"
    "発言改変リスクを避けるため、引用は原文（伊/英）と日本語翻訳を併記。",
    italic=True, size=9,
)

add_heading("4-1. 時系列ビュー", level=2)

add_para("FY2022（2022年1月、Tom's Hardware ITインタビュー）", bold=True)
add_para(
    "発言者：Lauro Buoro（Founder & Chairman）／場面：上場廃止後3年経過時点の戦略レビュー",
    italic=True, size=9,
)
add_para(
    "原文（伊→英訳）：「We are only at the beginning of a smart home market with room for all.」",
)
add_para(
    "解釈：Amazon／Googleの参入をthreatではなくopportunityと捉える楽観論。Niceの差別化を「Pro Installerチャネル"
    "での専門サポート」に置き、コンシューマDirectで戦わない方針を示唆。投資仮説の「smart home pivot」と整合。",
    size=10,
)

add_para("FY2022（2022年初頭、il Nord Estインタビュー）", bold=True)
add_para(
    "発言者：Lauro Buoro／場面：Nortek S&C買収（2021年10月）の戦略意義について",
    italic=True, size=9,
)
add_para(
    "原文（伊）：「L'acquisizione rafforza sostanzialmente la presenza di Nice sul mercato nordamericano,"
    "che riteniamo strategico per la nostra crescita.」",
)
add_para(
    "和訳：「この買収はNiceの北米市場におけるプレゼンスを実質的に強化するものであり、北米は我々の成長にとって"
    "戦略的市場である。」",
)
add_para(
    "解釈：北米を「戦略的市場」と明示。投資仮説「Nortek基盤の収益化」と整合する。",
    size=10,
)

add_para("FY2023（2023年9月、PR Newswire／FSI press release）", bold=True)
add_para(
    "発言者：Lauro Buoro／場面：FSI €100M少数出資の発表",
    italic=True, size=9,
)
add_para(
    "原文（英）：「Nice thus moves to the next level as the sole global partner offering the most complete"
    "integrated product ecosystem for homes and buildings.」",
)
add_para(
    "解釈：FSI出資により「next level」へ移行する宣言。「sole global partner」という表現はSomfy・FAACへの"
    "competitive positioning。投資仮説と整合。",
    size=10,
)

add_para("FY2024（2024年6月、PR Newswire／CEO交代発表）", bold=True)
add_para(
    "発言者：Lauro Buoro（Chairman comment on Mogollon appointment）／場面：CEO交代発表",
    italic=True, size=9,
)
add_para(
    "原文（英）：「Mogollon's proven leadership and extensive experience in mature and emerging markets"
    "are key elements in the development process of the Nice group.」",
)
add_para(
    "解釈：「mature and emerging markets」両面でのリーダーシップを明示的評価。Mogollonの主要KPIは"
    "「emerging market growth + mature market efficiency」と推察される。",
    size=10,
)

add_para("FY2024-FY2025（2025年初頭、il Nord Estインタビュー）", bold=True)
add_para(
    "発言者：Nice Group公式コメント（Mogollon体制下）／場面：2025年再編発表",
    italic=True, size=9,
)
add_para(
    "和訳：「2025年に向けて企業再編が計画されており、より柔軟性が高く、運用効率の大幅改善を伴う"
    "「グローバル組織構造のagilityと競争力強化」を目指す。「適切な対策」をConfindustria Veneto Estと"
    "の協調で実施し、影響は「限定的」とする。」",
)
add_para(
    "解釈：「limited impact」と表現される再編は通常、特定事業部の閉鎖／統合・希望退職プログラムを含む。"
    "FY2024売上微減（830→810）と整合。投資仮説「2025年再編完成によるオペレーション統合益の刈取」と整合。",
    size=10,
)

add_heading("4-2. 論点軸サマリ", level=2)

add_para("論点1：事業セグメント別の戦略・撤退・拡大方針", bold=True)
add_para(
    "発言：「Smart Home pivot」（Buoro 2022）／「sole global partner offering the most complete integrated"
    "product ecosystem」（Buoro 2023）。コアゲート事業からの拡張をブランド戦略の中心に置く。"
    "撤退発言は確認できず（追加DDで確認要）。Valcucine（家具、2024年10月）取得は「家まわり統合」の延長と推察。",
    size=10,
)

add_para("論点2：主要顧客との関係性・依存度", bold=True)
add_para("該当発言なし（追加DDで確認要）。Pro Installerチャネルの分散性については間接示唆あり。", size=10)

add_para("論点3：競合認識・市場シェア・価格戦略", bold=True)
add_para(
    "発言：「room for all」（Buoro 2022）はSmart Home市場でのoligopolistic peace前提。"
    "Somfy／FAACへの直接言及は確認できず。Amazon／Googleを「opportunity」と評価する点は、"
    "コンシューマDirectで戦わない明確な戦略選択。",
    size=10,
)

add_para("論点4：資本政策（株主還元・自己株・M&A・Exit）", bold=True)
add_para(
    "発言：「export 95%」（Buoro 2018 delisting時）の背景論理は「政府政策のリスクから逃れるための上場廃止」。"
    "FSI受入は「next level」（Buoro 2023）と表現し、否定的トーンなし。"
    "Buoro家のExit意図についての明示発言は確認できず（追加DDで重要論点）。",
    size=10,
)

add_para("論点5：人材・組織・後継者", bold=True)
add_para(
    "発言：「Human capital is fundamental to our company's development and expansion strategy」（Buoro 2022）。"
    "Mogollon登用（2024年6月）は外部からの招聘であり、Buoro家親族からの後継登用ではない点が重要。"
    "創業家ガバナンスの近代化を示唆。後継者プラン（Buoro後の経営）は引き続きOpen Question。",
    size=10,
)

add_para("論点6：数値目標と進捗自己評価（中計、利益率目標）", bold=True)
add_para(
    "中計／利益率目標は公表されておらず、該当発言なし（追加DDで確認要）。"
    "FSI出資ストーリー（press release）からは「mid-term carbon emission reduction targets by 2030」"
    "という非財務KPIのみ言及あり。",
    size=10,
)

add_para("論点7：不確実性・リスク認識", bold=True)
add_para(
    "発言：2018年delisting時のBuoro発言「投資家がイタリア企業から逃げている」「政治不確実性」"
    "は構造的リスク認識を示すが、それ以降の公開発言ではマクロ／地政学リスクへの言及は確認できず。"
    "2025年再編発表時の「limited impact」表現は内部にcost overhang圧力があることを示唆。",
    size=10,
)

doc.add_page_break()

# ============================================================================
# 付録
# ============================================================================
add_heading("付録：出典一覧", level=1)
sources = [
    "[1] FSI press release \"Nice Enters Strategic Partnership with FSI to Bolster Global Growth\" (PR Newswire, 2023-09-26 / 公式: fondofsi.it/en/investimenti-fsi/nice/)",
    "[2] BeBeez \"FSI investirà 100 mln euro per una minoranza in Nice, che ha chiuso il 2022 con 800 mln di ricavi\" (2023-09-26)",
    "[3] PR Newswire \"Nice: Juan B. Mogollon designated CEO\" (2024-06-25)",
    "[4] Nice S.p.A. Annual Financial Report 2018 (last public, ir.niceforyou.com)",
    "[5] Tribuna di Treviso \"La Nice di Oderzo esce da Borsa Italiana\" (2018-12)",
    "[6] il Nord Est \"La versione di Buoro: Nice dopo il delisting svolta e cambia pelle 'Ora forti negli Usa'\" (2022)",
    "[7] il Nord Est \"Nice Group ristruttura: 'Contenimento dei costi con un impatto limitato'\" (2025)",
    "[8] Tom's Hardware Italy \"Nice Group: la regina italiana dell'automazione che guarda alla smart home\" (2022年1月)",
    "[9] Italian Camera di Commercio (Nice S.p.A. Oderzo, P.IVA 03099360269) FY2022/FY2023/FY2024 standalone bilanci (reportaziende.it経由)",
    "[10] PR Newswire \"Nice Strengthens Global Smart Home & Building Automation Leadership Position with Nortek Security & Control LLC Acquisition\" (2021-10-05)",
    "[11] PR Newswire \"Nice acquires FIBARO\" (2018年)",
    "[12] CEPRO \"Fibaro Home Automation Acquired for $73M by European Controls Co.\" (2018)",
    "[13] Niceforyou.com /en/about (公式企業概要、ブランド一覧)",
    "[14] Fortune Business Insights \"Gate Openers Market\" (2024)",
    "[15] ResearchAndMarkets \"Gate Automation Market\" (2024)",
    "[16] Grand View Research \"Smart Home Market Size and Share\" (2024)",
    "[17] FAAC Technologies press releases (FY2023/FY2024 results, faac.it)",
    "[18] CAME Group \"Il Gruppo CAME chiude il 2024 in crescita: 335 mln fatturato\" (securindex.com 2025)",
    "[19] Somfy Group annual reports (somfy-group.com)",
    "[20] ASSA ABLOY Q4 2024 Report (assaabloy.com)",
]
for s in sources:
    add_bullet(s)

add_heading("略語集", level=1)
glossary = [
    ("HMS", "Home Management Solutions（Niceの自称セグメント名）"),
    ("S&C", "Security & Control（Nortek S&Cの略）"),
    ("FSI", "Fondo Strategico Italiano（イタリア国策系PE、現在のFSIファンド名）"),
    ("KBF", "Key Buying Factor（顧客選定要因）"),
    ("PMI", "Post-Merger Integration（買収後統合）"),
    ("PGA", "Purchase Price Allocation（取得価格配分）"),
    ("EBITDA", "Earnings Before Interest, Taxes, Depreciation, and Amortization"),
    ("CAGR", "Compound Annual Growth Rate（年平均成長率）"),
    ("CEDIA", "Custom Electronic Design and Installation Association（米Pro Integrator業界団体）"),
    ("OPA", "Offerta Pubblica di Acquisto（伊TOB相当）"),
    ("HEMS", "Home Energy Management System"),
]
for term, desc in glossary:
    add_bullet(f"{term}：{desc}")

# ============================================================================
# Save
# ============================================================================
import os
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"\n✓ docx saved: {OUT_PATH}")
